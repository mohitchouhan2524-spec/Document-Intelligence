"""
ingestion/scraper.py:
Loads documents from disk (PDF, DOCX, TXT, HTML, CSV) or URLs.
Returns a list of Document objects ready for chunking.
"""
from __future__ import annotations
import hashlib
from pathlib import Path

from loguru import logger

from src.models import Document


def _make_id(content: str) -> str:
    return hashlib.sha256(content.encode()).hexdigest()[:16]


class DocumentScraper:
    """Unified loader for local files and URLs."""

    SUPPORTED = {".pdf", ".docx", ".txt", ".html", ".htm", ".csv", ".xlsx"}

    # ── Public loaders

    def load_file(self, path: str | Path) -> Document:
        path = Path(path)
        if path.suffix.lower() not in self.SUPPORTED:
            raise ValueError(
                f"Unsupported file type: '{path.suffix}'\n"
                f"Supported: {sorted(self.SUPPORTED)}"
            )
        logger.info(f"Loading: {path.name}")
        content = self._extract(path)
        if not content.strip():
            raise ValueError(f"File loaded but extracted content is empty: {path.name}")
        return Document(
            doc_id=_make_id(content),
            source=str(path),
            content=content,
            metadata={"filename": path.name, "filetype": path.suffix.lstrip(".")},
        )

    def load_directory(self, directory: str | Path) -> list[Document]:
        directory = Path(directory)

        # ── Guard: directory must exist
        if not directory.exists():
            print(f"[ERROR] Directory not found: {directory.resolve()}")
            return []

        # ── Collect supported files (skip index/db/pkl files automatically) ───
        candidates = sorted(
            p for p in directory.rglob("*")
            if p.is_file() and p.suffix.lower() in self.SUPPORTED
        )

        if not candidates:
            all_files = [p for p in directory.rglob("*") if p.is_file()]
            exts      = sorted({p.suffix for p in all_files})
            print(f"[WARNING] No supported files found in: {directory.resolve()}")
            print(f"          Supported extensions : {sorted(self.SUPPORTED)}")
            print(f"          Extensions present   : {exts}")
            print(f"          Files present        :")
            for f in sorted(all_files):
                print(f"            {f.relative_to(directory)}")
            return []

        print(f"[INFO] Found {len(candidates)} supported file(s) in: {directory.resolve()}")

        # ── Load each file, report per-file result 
        docs:   list[Document]          = []
        failed: list[tuple[Path, Exception]] = []

        for path in candidates:
            try:
                doc = self.load_file(path)
                docs.append(doc)
                relative_path = path.relative_to(directory).as_posix()
                print(f"  ✓  {relative_path:<50}  {len(doc.content):>7} chars")
            except Exception as e:
                failed.append((path, e))
                relative_path = path.relative_to(directory).as_posix()
                print(f"  ✗  {relative_path:<50}  {type(e).__name__}: {e}")
                logger.warning(f"Failed to load {path}: {e}")

        # ── Summary
        print(f"\n[INFO] Result: {len(docs)} loaded, {len(failed)} failed "
              f"(total {len(candidates)} files)")

        if failed:
            print("\n[HINT] Common causes for failures:")
            print("  PDF parse error  → pip install pypdf")
            print("  DOCX parse error → pip install python-docx")
            print("  Scanned PDF      → pip install unstructured pytesseract")
            print("  Encoding error   → file may be corrupt or password-protected")

        logger.info(f"load_directory: {len(docs)} loaded, {len(failed)} failed from {directory}")
        return docs

    def load_url(self, url: str) -> Document:
        import httpx
        from bs4 import BeautifulSoup

        logger.info(f"Fetching URL: {url}")
        response = httpx.get(url, follow_redirects=True, timeout=30)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        content = soup.get_text(separator="\n", strip=True)
        return Document(
            doc_id=_make_id(content),
            source=url,
            content=content,
            metadata={"url": url, "filetype": "html"},
        )

    # ── Private extraction 

    def _extract(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._pdf(path)
        if suffix == ".docx":
            return self._docx(path)
        if suffix in (".html", ".htm"):
            return self._html(path)
        if suffix in (".txt", ".csv"):
            return path.read_text(encoding="utf-8", errors="ignore")
        # fallback: unstructured handles xlsx and edge cases
        return self._unstructured(path)

    def _pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required to read PDF files.\n"
                "Install it with: pip install pypdf"
            )
        reader = PdfReader(str(path))
        if not reader.pages:
            raise ValueError(f"PDF has no pages: {path.name}")

        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
            else:
                # Scanned page — try OCR via unstructured
                logger.debug(f"Page {i} of {path.name} is empty, trying OCR")
                try:
                    ocr_text = self._unstructured(path)
                    if ocr_text.strip():
                        return ocr_text   # unstructured parsed the whole PDF
                except Exception as ocr_err:
                    logger.warning(f"OCR failed for {path.name}: {ocr_err}")

        if not pages:
            raise ValueError(
                f"No extractable text in {path.name}. "
                "The PDF may be scanned — install unstructured + pytesseract for OCR."
            )
        return "\n".join(pages)

    def _docx(self, path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required to read DOCX files.\n"
                "Install it with: pip install python-docx"
            )
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _html(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required to read HTML files.\n"
                "Install it with: pip install beautifulsoup4 lxml"
            )
        soup = BeautifulSoup(path.read_text(encoding="utf-8"), "lxml")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    def _unstructured(self, path: Path) -> str:
        try:
            from unstructured.partition.auto import partition
        except ImportError:
            raise ImportError(
                "unstructured is required for this file type.\n"
                "Install it with: pip install unstructured"
            )
        elements = partition(filename=str(path))
        return "\n".join(str(e) for e in elements)